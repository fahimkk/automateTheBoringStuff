import docx

doc = docx.Document()
doc.add_paragraph('Hello World!','Title')

paraObj1 = doc.add_paragraph('This is a second paragraph.')
paraObj2 = doc.add_paragraph('This is add as a 3rd paragraph')

paraObj1.add_run('this is aded for 2nd para as run.').underline = True

doc.save('multipleParagraphs.docx')




