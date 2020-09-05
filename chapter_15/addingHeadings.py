import docx

doc = docx.Document()

doc.add_heading('Pyhon Documents ', 1)
doc.add_heading('Header 2', 2)
doc.add_heading('work 3', 3)
doc.add_heading('content 4', 4)

doc.save('headings.docx')
