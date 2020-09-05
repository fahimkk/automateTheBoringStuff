import docx
import readDocx

doc_file = docx.Document('demo.docx')
print(len(doc_file.paragraphs))

for paraNum in range(len(doc_file.paragraphs)):
    para = doc_file.paragraphs[paraNum]
    print(para.text, end=' : ')
    print(len(para.runs),end=' : ')
    for i in range(len(para.runs)):
        print(para.runs[i].text, end=',')
    print()

print()
print(doc_file.paragraphs[1].text)
print(len(doc_file.paragraphs[1].runs))
print()

print(readDocx.getText('demo.docx'))