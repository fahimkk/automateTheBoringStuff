import docx
import logging

logging.basicConfig(level=logging.DEBUG)

doc = docx.Document('demo.docx')
logging.debug(doc.paragraphs[0].text)

doc.paragraphs[0].style = 'Normal'

doc.paragraphs[1].runs[0].style = 'QuoteChar'
logging.debug(doc.paragraphs[1].runs[0].text)
doc.paragraphs[1].runs[1].underline = True
logging.debug(doc.paragraphs[1].runs[1].text)
doc.paragraphs[1].runs[2].underline = True
logging.debug(doc.paragraphs[1].runs[2].text)

doc.save('restyled.docx')


